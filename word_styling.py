from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.oxml import OxmlElement

def export_doc(template_path, output_path, daten):
    doc = Document(template_path)

    def ersetze(text):
        return text if not text else text.format(**daten)

    for p in doc.paragraphs:
        p.text = ersetze(p.text)

    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                cell.text = ersetze(cell.text)

    for para in doc.paragraphs:
        if "{Kundenname}" in para.text:
            for run in para.runs:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.bold = True
            p = para._element
            pPr = p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '8')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '000000')
            pBdr.append(bottom)
            pPr.append(pBdr)

    kleiner_text = "Diese Übersicht wurde automatisch vom Reisebüro Hülsmann"
    for para in doc.paragraphs:
        if kleiner_text in para.text:
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(120, 120, 120)
                run.italic = True

    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        if len(table.columns) >= 3:
            table.columns[0].width = Inches(2.8)
            table.columns[1].width = Inches(1.6)
            table.columns[2].width = Inches(1.6)

        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                cell.width = table.columns[col_index].width
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(11)
                        run.font.name = 'Arial'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
                        if row_index == 0 or col_index == 0:
                            run.bold = True

                if row_index == 0 or col_index == 0:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), "eeeeee")
                    cell._tc.get_or_add_tcPr().append(shading)

    ueberschriften = [
        "Welche Versicherungen sind wichtig",
        "Abschlussfristen",
        "Reiserücktritts-Versicherung",
        "Reisekranken-Versicherung",
        "RundumSorglos-Schutz"
    ]
    for para in doc.paragraphs:
        if para.text.strip() in ueberschriften:
            for run in para.runs:
                run.bold = True

    doc.save(output_path)
    return output_path