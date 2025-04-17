
import os
import re
import textwrap
SYNONYME = {
    "selbstbeteiligung": ["selbstbehalt", "eigenanteil", "sb"],
    "reiserücktritt": ["rücktritt", "stornierung"],
    "versicherung": ["schutz", "tarif", "deckung"],
    "krankheit": ["corona", "covid", "erkrankung"],
    "rundumsorglos": ["paket", "komplettschutz"]
}
import streamlit as st
import pandas as pd
import fitz  # PyMuPDF
from openai import OpenAI
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

def pdf_suche(pfad, suchbegriff):
    doc = fitz.open(pfad)
    treffer = []
    suchbegriffe = [suchbegriff.lower()] + SYNONYME.get(suchbegriff.lower(), [])
    
    for seite in doc:
        text = seite.get_text()
        text_lower = text.lower()
        if any(sb in text_lower for sb in suchbegriffe):
            auszug = re.sub(r"\s+", " ", text.strip())
            treffer.append((seite.number + 1, auszug[:1000]))  # max 1000 Zeichen
    return treffer

def highlight(text, wort):
    clean_text = re.sub(r"[�•⍰]", " ", text)
    pattern = re.compile(f"(?i)({re.escape(wort)})")
    return pattern.sub(r"<mark>\1</mark>", clean_text)
    
from datetime import datetime, date
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.oxml import OxmlElement

st.set_page_config(page_title="Der Ergo Chuck", page_icon="🦾", layout="centered")
st.image("SmallLogoBW_png.png", width=250)

st.markdown(
    "<div style='font-size:18px; color:#111; font-style:italic; margin-bottom:1em;'>"
    "Wenn Chuck Norris eine Reise plant, versichert sich das Zielland."
    "</div>",
    unsafe_allow_html=True
)

st.markdown(
    "<h2 style='font-size:28px; color:#111;'>Der Ergo Chuck – Berechnung & Angebot</h2>",
    unsafe_allow_html=True
)

def parse_datum(text):
    try:
        if "." in text:
            return datetime.strptime(text.strip(), "%d.%m.%Y").date()
        elif len(text.strip()) == 4:
            return date.today().replace(day=int(text[:2]), month=int(text[2:]))
        elif len(text.strip()) == 8:
            return datetime.strptime(text.strip(), "%d%m%Y").date()
    except:
        return None
        
def ermittle_zielgebiet(code):
    zielgebiet = st.radio("Zielgebiet", ["Europa", "Welt"], index=0)
    return zielgebiet

def ermittle_altersgruppe(alter):
    if alter <= 40: return "bis 40 Jahre"
    elif 41 <= alter <= 64: return "41–64 Jahre"
    return "ab 65 Jahre"

def ermittle_personengruppe(alter_liste):
    if len(alter_liste) == 1: return "Einzelperson"
    elif len(alter_liste) == 2: return "Paar"
    return "Familie"

def berechne_reisetage(von, bis):
    return max(1, (bis - von).days + 1)

def fmt(p, c, preis):
    p = float(p)
    betrag = p * preis if p < 1.0 else p
    return f"{betrag:.2f}".replace(".", ",") + f" € ({c})"

def first_hit(df):
    col = [c for c in df.columns if c.strip().lower() == "reisepreis bis"]
    if col:
        return df.sort_values(col[0]).iloc[0]
    return df.iloc[0]

def bereinige_export_daten(daten):
    export_daten = {}
    for key, value in daten.items():
        if isinstance(value, str) and "€" in value and "(" in value and ")" in value:
            betrag = value.split("€")[0].strip()
            export_daten[key] = betrag + " €"
        else:
            export_daten[key] = value
    return export_daten

def export_doc(template_path, output_path, daten):
    doc = Document(template_path)

    def ersetze(text): return text if not text else text.format(**daten)
    for p in doc.paragraphs:
        p.text = ersetze(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                cell.text = ersetze(cell.text)

    for para in doc.paragraphs:
        if "{Kundenname}" in para.text:
            for run in para.runs:
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.bold = True
            p = para._element
            pPr = p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '8')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '000000')
            pBdr.append(bottom)
            pPr.append(pBdr)
        
    kleiner_text = "Diese Übersicht wurde automatisch vom Reisebüro Hülsmann"
    for para in doc.paragraphs:
        if kleiner_text in para.text:
            for run in para.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(120, 120, 120)
                run.italic = True
    
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    ueberschriften = [
        "Welche Versicherungen sind wichtig",
        "Abschlussfristen",
        "Reiserücktritts-Versicherung",
        "Reisekranken-Versicherung",
        "RundumSorglos-Schutz"
    ]
    for para in doc.paragraphs:
        if para.text.strip() in ueberschriften:
            for run in para.runs:
                run.bold = True

    doc.save(output_path)
    return output_path

with st.form("eingabeformular"):
    name = st.text_input("Kundenname")
    zielgebiet = st.radio("Zielgebiet", ["Europa", "Welt"], index=0)
    preis = st.number_input("Reisepreis (€)", min_value=0.0)
    alter_text = st.text_input("Alter (z. B. 45 48)")
    von_raw = st.text_input("Reise von (TTMM oder TT.MM.JJJJ)")
    bis_raw = st.text_input("Reise bis (TTMM oder TT.MM.JJJJ)")
    submit = st.form_submit_button("Tarife anzeigen")

if submit:
    try:
        von = parse_datum(von_raw)
        bis = parse_datum(bis_raw)
        if not von or not bis:
            st.error("❌ Bitte gültige Datumswerte eingeben.")
            st.stop()

        alter_liste = [int(a) for a in alter_text.strip().split()]
        max_alter = max(alter_liste)
        altersgruppe = ermittle_altersgruppe(max_alter)
        personengruppe = ermittle_personengruppe(alter_liste)
        reisetage = berechne_reisetage(von, bis)

        excel = pd.ExcelFile("ergo.xlsx")
        blatt = {
            "rrv_ew_mit": "rrv-ev-mit", "rrv_ew_ohne": "rrv-ev-ohne",
            "rrv_jw_mit": "rrv-jv-mit", "rrv_jw_ohne": "rrv-jv-ohne",
            "rrv_jw_spf_mit": "rrv-jv-spf-mit", "rrv_jw_spf_ohne": "rrv-jv-spf-ohne",
            "kv_ew_mit": "kv-ev-mit", "kv_ew_ohne": "kv-ev-ohne",
            "kv_jw_mit": "kv-jv-mit", "kv_jw_ohne": "kv-jv-ohne",
            "rus_ew_mit": "rus-ev-mit", "rus_ew_ohne": "rus-ev-ohne",
            "rus_jw_mit": "rus-jv-mit", "rus_jw_ohne": "rus-jv-ohne"
        }
        t = {key: excel.parse(name) for key, name in blatt.items()}

        def tarif(d, *args): return fmt(*args, preis) if not d.empty else "–"
        def first(d, cond, keys): f = d[cond]; return tarif(f, *first_hit(f)[keys]) if not f.empty else "–"

        def kv_einmal(name):
            d = t[name]
            f = d[
                (d["Zielgebiet"].str.strip().str.lower() == zielgebiet.lower()) &
                (d["Altersgruppe"].str.strip() == altersgruppe) &
                (d["Personengruppe"].str.strip().str.lower() == personengruppe.lower())
            ]
            if not f.empty:
                tagespreis = float(f.iloc[0]["Tagesprämie"])
                return fmt(round(reisetage * tagespreis, 2), f.iloc[0]["Tarifcode"], preis)
            return "–"

        
        daten = {
            "Kundenname": name,
            "Reisedatum": f"{von.strftime('%d.%m.%Y')} – {bis.strftime('%d.%m.%Y')}",
            "Reisepreis": f"{preis:,.2f} €".replace(".", ","),
            "Anzahl": str(len(alter_liste)),
            "Alter": ", ".join(str(a) for a in alter_liste),
            "Reiseziel": zielgebiet,
            "Reiseruecktritt_Einmal_mit_SB": first(t["rrv_ew_mit"], t["rrv_ew_mit"]["Reisepreis bis"] >= preis, ["Prämie", "Tarifcode"]),
            "Reiseruecktritt_Einmal_ohne_SB": first(t["rrv_ew_ohne"], (t["rrv_ew_ohne"]["Reisepreis bis"] >= preis) & (t["rrv_ew_ohne"]["Altersgruppe"].str.strip() == altersgruppe), ["Prämie", "Tarifcode"]),
            "Reiseruecktritt_Jahres_mit_SB": first(t["rrv_jw_mit"], (t["rrv_jw_mit"]["Reisepreis bis"] >= preis) & (t["rrv_jw_mit"]["Altersgruppe"].str.strip() == altersgruppe) & (t["rrv_jw_mit"]["Personengruppe"].str.lower().str.strip() == personengruppe.lower()), ["Prämie", "Tarifcode"]),
            "Reiseruecktritt_Jahres_ohne_SB": first(t["rrv_jw_ohne"], (t["rrv_jw_ohne"]["Reisepreis bis"] >= preis) & (t["rrv_jw_ohne"]["Altersgruppe"].str.strip() == altersgruppe) & (t["rrv_jw_ohne"]["Personengruppe"].str.lower().str.strip() == personengruppe.lower()), ["Prämie", "Tarifcode"]),
            "Reiseruecktritt_Jahres_Sparfuchs_mit_SB": first(t["rrv_jw_spf_mit"], (t["rrv_jw_spf_mit"]["Reisepreis bis"] >= preis) & (t["rrv_jw_spf_mit"]["Altersgruppe"].str.strip() == altersgruppe) & (t["rrv_jw_spf_mit"]["Personengruppe"].str.lower().str.strip() == personengruppe.lower()), ["Prämie", "Tarifcode"]),
            "Reiseruecktritt_Jahres_Sparfuchs_ohne_SB": first(t["rrv_jw_spf_ohne"], (t["rrv_jw_spf_ohne"]["Reisepreis bis"] >= preis) & (t["rrv_jw_spf_ohne"]["Altersgruppe"].str.strip() == altersgruppe) & (t["rrv_jw_spf_ohne"]["Personengruppe"].str.lower().str.strip() == personengruppe.lower()), ["Prämie", "Tarifcode"]),
            "Reisekranken_Einmal_mit_SB": kv_einmal("kv_ew_mit"),
            "Reisekranken_Einmal_ohne_SB": kv_einmal("kv_ew_ohne"),
            "Reisekranken_Jahres_mit_SB": first(t["kv_jw_mit"], (t["kv_jw_mit"]["Altersgruppe"].str.strip() == altersgruppe) & (t["kv_jw_mit"]["Personengruppe"].str.lower().str.strip() == personengruppe.lower()), ["Prämie", "Tarifcode"]),
            "Reisekranken_Jahres_ohne_SB": first(t["kv_jw_ohne"], (t["kv_jw_ohne"]["Altersgruppe"].str.strip() == altersgruppe) & (t["kv_jw_ohne"]["Personengruppe"].str.lower().str.strip() == personengruppe.lower()), ["Prämie", "Tarifcode"]),
            "RundumSorglos_Einmal_mit_SB": first(t["rus_ew_mit"], (t["rus_ew_mit"]["Reisepreis bis"] >= preis) & (t["rus_ew_mit"]["Zielgebiet"].str.lower().str.strip() == zielgebiet.lower()), ["Prämie", "Tarifcode"]),
            "RundumSorglos_Einmal_ohne_SB": first(t["rus_ew_ohne"], (t["rus_ew_ohne"]["Reisepreis bis"] >= preis) & (t["rus_ew_ohne"]["Zielgebiet"].str.lower().str.strip() == zielgebiet.lower()) & (t["rus_ew_ohne"]["Altersgruppe"].str.strip() == altersgruppe), ["Prämie", "Tarifcode"]),
            "RundumSorglos_Jahres_mit_SB": first(t["rus_jw_mit"], (t["rus_jw_mit"]["Reisepreis bis"] >= preis) & (t["rus_jw_mit"]["Altersgruppe"].str.strip() == altersgruppe) & (t["rus_jw_mit"]["Personengruppe"].str.lower().str.strip() == personengruppe.lower()), ["Prämie", "Tarifcode"]),
            "RundumSorglos_Jahres_ohne_SB": first(t["rus_jw_ohne"], (t["rus_jw_ohne"]["Reisepreis bis"] >= preis) & (t["rus_jw_ohne"]["Altersgruppe"].str.strip() == altersgruppe) & (t["rus_jw_ohne"]["Personengruppe"].str.lower().str.strip() == personengruppe.lower()), ["Prämie", "Tarifcode"]),
        }
        st.session_state["word_daten"] = daten
            
        df = pd.DataFrame([
            ["Reiserücktritt", "Einmal", daten["Reiseruecktritt_Einmal_mit_SB"], daten["Reiseruecktritt_Einmal_ohne_SB"]],
            ["", "Jahres", daten["Reiseruecktritt_Jahres_mit_SB"], daten["Reiseruecktritt_Jahres_ohne_SB"]],
            ["", "Sparfuchs", daten["Reiseruecktritt_Jahres_Sparfuchs_mit_SB"], daten["Reiseruecktritt_Jahres_Sparfuchs_ohne_SB"]],
            ["Reisekranken", "Einmal", daten["Reisekranken_Einmal_mit_SB"], daten["Reisekranken_Einmal_ohne_SB"]],
            ["", "Jahres", daten["Reisekranken_Jahres_mit_SB"], daten["Reisekranken_Jahres_ohne_SB"]],
            ["RundumSorglos", "Einmal", daten["RundumSorglos_Einmal_mit_SB"], daten["RundumSorglos_Einmal_ohne_SB"]],
            ["", "Jahres", daten["RundumSorglos_Jahres_mit_SB"], daten["RundumSorglos_Jahres_ohne_SB"]],
        ], columns=["Produktgruppe", "Tarif", "mit SB", "ohne SB"])
        st.subheader("📊 Gruppierte Tarifübersicht")
        st.table(df)


    except Exception as e:
        st.error(f"❌ Fehler: {e}")

if "word_daten" in st.session_state:
    st.subheader("📄 Word-Angebot")
    if st.button("📄 Word-Angebot erstellen"):
        if not os.path.exists("angebot.docx"):
            st.warning("⚠️ Datei 'angebot.docx' fehlt!")
        else:
            saubere_daten = bereinige_export_daten(st.session_state["word_daten"])
            file_path = "angebot_fertig_streamlit.docx"
            export_doc("angebot.docx", file_path, saubere_daten)
            with open(file_path, "rb") as f:
                file_bytes = f.read()
            st.download_button("📥 Angebot herunterladen", file_bytes, file_name=file_path)

st.subheader("🤖 Beratung zur ERGO-Reiseversicherung")

frage_gpt = st.text_input("Welche Frage haben Sie zur Versicherung?", placeholder="z. B. Was ist bei Corona versichert?")

if frage_gpt.strip():
    with st.spinner("Durchsuche PDF und frage GPT …"):
        fundstellen = pdf_suche("ergo_tarife.pdf", frage_gpt)

        if not fundstellen:
            st.warning("📄 Keine passenden Textstellen in der PDF gefunden.")
        else:
            # Kontext aus max. 2 Fundstellen
            kontext = "\n\n".join([f"Seite {s}:\n{t}" for s, t in fundstellen[:2]])

            # 📄 Fundstellen anzeigen (max. 3) – NUR wenn vorhanden
            with st.expander("📄 Gefundene Textstellen anzeigen"):
                for i, (s, t) in enumerate(fundstellen[:3], 1):
                    st.markdown(f"**{i}. Seite {s}**")
                    st.markdown(textwrap.shorten(t, width=600, placeholder=" …"), unsafe_allow_html=True)

            # 🤖 GPT-System-Prompt mit Synonymerkennung
            system_prompt = (
                "Du bist ein digitaler Versicherungsberater für Reisebüro Hülsmann. "
                "Beantworte ausschließlich Fragen zu Reiserücktritts-, Reisekranken- oder RundumSorglos-Versicherungen "
                "auf Grundlage der folgenden PDF-Auszüge.\n\n"
                "Berücksichtige bei der Interpretation auch Begriffe mit ähnlicher Bedeutung. "
                "Zum Beispiel:\n"
                "- Selbstbeteiligung ≈ Selbstbehalt ≈ SB ≈ Eigenanteil\n"
                "- Reiserücktritt ≈ Rücktritt ≈ Stornierung\n"
                "- Krankheit ≈ Corona ≈ COVID ≈ Quarantäne\n\n"
                "Wenn du keine ausreichende Information findest, sage bitte klar: "
                "'Dazu liegt mir keine Information vor.'"
            )

            # 🧠 GPT-Request mit OpenAI v1
            response = client.chat.completions.create(
                model="gpt-4-turbo",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"Frage: {frage_gpt}\n\nPDF-Auszüge:\n{kontext}"}
                ],
                temperature=0.3
            )

            antwort = response.choices[0].message.content
            st.success(antwort)
